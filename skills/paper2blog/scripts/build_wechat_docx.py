#!/usr/bin/env python3
"""Build a simple editor-friendly Paper2Blog DOCX from a JSON outline.

Usage:
    python scripts/build_wechat_docx.py outline.json output.docx [--lang zh|en]

The JSON file should contain:
{
  "title": "...",
  "subtitle": "...",
  "lang": "zh",            # optional: "zh" (微软雅黑 + Arial) or "en" (Arial)
  "fonts": {               # optional per-document font override
    "latin": "Arial",
    "east_asia": "Microsoft YaHei"
  },
  "blocks": [
    {"type": "paragraph", "text": "..."},
    {"type": "heading", "level": 1, "text": "..."},
    {"type": "figure", "path": "assets/figure.png", "caption": "图1 ...", "width_inches": 6.2},
    {"type": "table", "caption": "表1 ...", "headers": ["方法", "结果"], "rows": [["A", "1"]]}
  ]
}

Fonts are chosen per language:
  - zh -> Microsoft YaHei (微软雅黑) for Chinese glyphs, Arial for Latin text.
  - en -> Arial for Latin text, Microsoft YaHei as CJK fallback.
If --lang is omitted, the language is taken from the outline "lang" field, or
inferred from the output filename (blog_zh.docx / blog_en.docx).
Manual font overrides can be supplied with the outline "fonts" field or CLI
flags. CLI font flags take precedence over outline fonts.
"""

from __future__ import annotations

import argparse
import json
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


try:
    from PIL import Image
except Exception:  # pragma: no cover - optional dependency
    Image = None


CONTENT_WIDTH_DXA = 9360

# Per-language font profiles.
#   latin     -> applied as the run/style font name (w:ascii / w:hAnsi)
#   east_asia -> applied as the w:eastAsia font for CJK glyphs
# zh: the Chinese WeChat article renders Chinese glyphs in 微软雅黑 (Microsoft
#     YaHei) and Latin text (embedded English terms, numbers) in Arial.
# en: the English research blog renders Latin text in Arial (a common
#     sans-serif face); 微软雅黑 is kept only as a CJK fallback for any stray
#     Chinese characters.
FONT_PROFILES = {
    "zh": {"latin": "Arial", "east_asia": "Microsoft YaHei"},
    "en": {"latin": "Arial", "east_asia": "Microsoft YaHei"},
}
DEFAULT_LANG = "zh"

# Active profile; overridden by build() via select_fonts() before any run/style
# is created. Defaults to the Chinese profile so the script is safe to import.
_FONTS = FONT_PROFILES[DEFAULT_LANG]


def select_fonts(lang: str, overrides: dict | None = None) -> dict:
    """Set the module-level font profile for the given language.

    `overrides` may contain "latin" and/or "east_asia" font names. This keeps
    the default language profiles stable while allowing a specific article to
    match an editor's font request without editing this script.
    """
    global _FONTS
    profile = dict(FONT_PROFILES.get(lang, FONT_PROFILES[DEFAULT_LANG]))
    if isinstance(overrides, dict):
        for key in ("latin", "east_asia"):
            value = overrides.get(key)
            if isinstance(value, str) and value.strip():
                profile[key] = value.strip()
    _FONTS = profile
    return _FONTS


def infer_lang(output_path: Path) -> str:
    """Best-effort language guess from the output filename (blog_zh / blog_en)."""
    stem = output_path.stem.lower()
    if stem.endswith("_en") or "blog_en" in stem:
        return "en"
    if stem.endswith("_zh") or "blog_zh" in stem:
        return "zh"
    return DEFAULT_LANG


def apply_run_font(run, east_asia: str | None = None, latin: str | None = None) -> None:
    latin = latin or _FONTS["latin"]
    east_asia = east_asia or _FONTS["east_asia"]
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(style, size_pt: float, color: str = "000000", bold: bool = False) -> None:
    style.font.name = _FONTS["latin"]
    style._element.rPr.rFonts.set(qn("w:eastAsia"), _FONTS["east_asia"])
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def get_or_add_style(styles, name: str, style_type):
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, style_type)


def prepare_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    set_style_font(styles["Normal"], 11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    body = get_or_add_style(styles, "Wechat Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(body, 11)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.1
    body.paragraph_format.first_line_indent = Inches(0.28)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    caption = get_or_add_style(styles, "Wechat Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(caption, 9.5, "555555")
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.1

    subtitle = get_or_add_style(styles, "Wechat Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(subtitle, 13, "555555")
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.line_spacing = 1.1

    set_style_font(styles["Title"], 22, "0B2545", bold=True)
    set_style_font(styles["Heading 1"], 16, "2E74B5", bold=True)
    set_style_font(styles["Heading 2"], 13, "2E74B5", bold=True)
    set_style_font(styles["Heading 3"], 12, "1F4D78", bold=True)
    return doc


def add_text_paragraph(doc: Document, text: str, style_name: str = "Wechat Body") -> None:
    p = doc.add_paragraph(style=style_name)
    r = p.add_run(text)
    apply_run_font(r)


def add_centered_text(doc: Document, text: str, style_name: str, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style_name)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    apply_run_font(r)


def safe_image_path(path: Path, tmp_dir: Path) -> Path:
    if Image is None:
        return path
    im = Image.open(path)
    if im.mode not in ("RGBA", "LA"):
        return path
    rgba = im.convert("RGBA")
    bg = Image.new("RGB", rgba.size, "white")
    bg.paste(rgba, mask=rgba.getchannel("A"))
    out = tmp_dir / f"{path.stem}_white.png"
    bg.save(out, optimize=True)
    return out


def add_figure(doc: Document, base_dir: Path, block: dict, tmp_dir: Path) -> None:
    image_path = (base_dir / block["path"]).resolve()
    image_path = safe_image_path(image_path, tmp_dir)
    width = float(block.get("width_inches", 6.2))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Inches(width))
    caption = block.get("caption")
    if caption:
        add_centered_text(doc, caption, "Wechat Caption")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_width(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)


def add_table(doc: Document, block: dict) -> None:
    caption = block.get("caption")
    if caption:
        add_centered_text(doc, caption, "Wechat Caption", bold=True)

    headers = block["headers"]
    rows = block["rows"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = str(header)
        set_cell_shading(cell, "F2F4F7")

    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)

    col_width = CONTENT_WIDTH_DXA // len(headers)
    widths = [col_width] * len(headers)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_width(table, widths)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    apply_run_font(run)
                    run.font.size = Pt(9.2)
                    if row_idx == 0:
                        run.bold = True

    note = block.get("note")
    if note:
        add_text_paragraph(doc, note, "Wechat Caption")


def build(
    outline_path: Path,
    output_path: Path,
    lang: str | None = None,
    font_overrides: dict | None = None,
) -> None:
    outline = json.loads(outline_path.read_text(encoding="utf-8"))
    base_dir = outline_path.parent

    # Font language precedence: explicit --lang > outline "lang" field >
    # inference from the output filename (blog_zh / blog_en).
    resolved_lang = (lang or outline.get("lang") or infer_lang(output_path)).lower()
    if resolved_lang not in FONT_PROFILES:
        resolved_lang = DEFAULT_LANG

    outline_fonts = outline.get("fonts") if isinstance(outline.get("fonts"), dict) else {}
    merged_font_overrides = dict(outline_fonts)
    if isinstance(font_overrides, dict):
        for key, value in font_overrides.items():
            if value:
                merged_font_overrides[key] = value
    select_fonts(resolved_lang, merged_font_overrides)

    doc = prepare_doc()

    if outline.get("title"):
        add_centered_text(doc, outline["title"], "Title", bold=True)
    if outline.get("subtitle"):
        add_centered_text(doc, outline["subtitle"], "Wechat Subtitle")

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        for block in outline.get("blocks", []):
            kind = block.get("type")
            if kind == "paragraph":
                add_text_paragraph(doc, block["text"])
            elif kind == "heading":
                level = int(block.get("level", 1))
                doc.add_heading(block["text"], level=max(1, min(level, 3)))
            elif kind == "figure":
                add_figure(doc, base_dir, block, tmp_dir)
            elif kind == "table":
                add_table(doc, block)
            elif kind == "caption":
                add_centered_text(doc, block["text"], "Wechat Caption")
            else:
                raise ValueError(f"Unsupported block type: {kind}")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("outline", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument(
        "--lang",
        choices=sorted(FONT_PROFILES),
        default=None,
        help=(
            "Language font profile: 'zh' uses Microsoft YaHei (微软雅黑) for "
            "Chinese and Arial for Latin, 'en' uses Arial. If "
            "omitted, inferred from the "
            "outline 'lang' field or the output filename (blog_zh / blog_en)."
        ),
    )
    parser.add_argument(
        "--latin-font",
        default=None,
        help="Override the Latin/ascii font for this output document.",
    )
    parser.add_argument(
        "--east-asia-font",
        default=None,
        help="Override the East Asian/CJK font for this output document.",
    )
    args = parser.parse_args()
    build(
        args.outline.resolve(),
        args.output.resolve(),
        lang=args.lang,
        font_overrides={
            "latin": args.latin_font,
            "east_asia": args.east_asia_font,
        },
    )
    print(args.output.resolve())


if __name__ == "__main__":
    main()
